import json
import logging
from io import BytesIO, StringIO

from config import KNOWLEDGE_BASE_DIR
from database import db
from models import Document, DocumentChunk
from services.pinecone_service import delete_from_pinecone, list_all_vectors, upsert_to_pinecone
from utils.embeddings import chunk_text

# Configure logging
logger = logging.getLogger(__name__)
# Logging is already configured in app.py
logger.debug("File service logger configured with console output")

# Log knowledge base directory
logger.info(f"Knowledge base directory set to: {KNOWLEDGE_BASE_DIR}")


def _safe_preview(text: str, limit: int = 120) -> str:
    """Return a short preview for logs without leaking large content."""
    if not isinstance(text, str):
        return ""
    return (text[:limit] + "...") if len(text) > limit else text


def extract_text_and_metadata(file_type: str, raw: bytes | str) -> tuple[str, dict]:
    """
    Extract normalized text and optional metadata (like doc_type) from uploaded payloads.

    Supports: txt, md, json, jsonl, csv, docx, pdf, xlsx
    """
    try:
        detected_doc_type = None
        ft = (file_type or "").lower()

        # Normalize raw to text where appropriate; keep bytes for binary parsers
        raw_text: str | None = None
        if isinstance(raw, bytes):
            try:
                raw_text = raw.decode("utf-8")
            except Exception:
                raw_text = None
        else:
            raw_text = raw

        if ft in {"txt", "md"}:
            return raw_text or "", {}

        if ft == "json":
            try:
                obj = json.loads(raw_text or "")
                if isinstance(obj, dict):
                    detected_doc_type = (
                        obj.get("doc_type")
                        or obj.get("type")
                        or obj.get("Category")
                        or obj.get("Type")
                    )
                return json.dumps(obj, indent=2), {"doc_type": (detected_doc_type or None)}
            except Exception:
                return raw_text or "", {}

        if ft == "jsonl":
            lines = (raw_text or "").splitlines()
            records: list[str] = []
            for idx, line in enumerate(lines):
                try:
                    obj = json.loads(line)
                    if isinstance(obj, dict) and detected_doc_type is None:
                        detected_doc_type = (
                            obj.get("doc_type")
                            or obj.get("type")
                            or obj.get("Category")
                            or obj.get("Type")
                        )
                    records.append(f"Record {idx+1}: {json.dumps(obj, indent=2)}")
                except Exception:
                    records.append(f"Line {idx+1}: {line}")
            return "\n\n".join(records), {"doc_type": (detected_doc_type or None)}

        if ft == "csv":
            import csv

            sio = StringIO(raw_text or "")
            reader = csv.reader(sio)
            rows = list(reader)
            if not rows:
                return "", {}
            header = rows[0]
            # infer doc_type from known header names
            detected_idx = None
            for i, h in enumerate(header):
                if (h or "").strip().lower() in {"category", "doc_type", "type"}:
                    detected_idx = i
                    break
            if detected_idx is not None and len(rows) > 1 and len(rows[1]) > detected_idx:
                detected_doc_type = (rows[1][detected_idx] or "").strip()
            formatted: list[str] = ["CSV Headers: " + ", ".join(header)]
            for i, row in enumerate(rows[1:]):
                row_dict = {header[j]: value for j, value in enumerate(row) if j < len(header)}
                formatted.append(f"Row {i+1}: {json.dumps(row_dict, indent=2)}")
            return "\n\n".join(formatted), {"doc_type": (detected_doc_type or None)}

        if ft == "xlsx":
            try:
                from openpyxl import load_workbook

                wb = load_workbook(
                    filename=BytesIO(
                        raw if isinstance(raw, bytes) else (raw_text or "").encode("utf-8")
                    ),
                    read_only=True,
                    data_only=True,
                )
                output: list[str] = []
                for ws in wb.worksheets:
                    output.append(f"# Sheet: {ws.title}")
                    rows = list(ws.iter_rows(values_only=True))
                    if not rows:
                        continue
                    header = [str(h) if h is not None else "" for h in rows[0]]
                    output.append("XLSX Headers: " + ", ".join(header))
                    for i, h in enumerate(header):
                        if h.strip().lower() in {"category", "doc_type", "type"} and len(rows) > 1:
                            first_row = rows[1]
                            if i < len(first_row) and first_row[i] is not None:
                                detected_doc_type = str(first_row[i]).strip()
                                break
                    for i, row in enumerate(rows[1:], start=1):
                        row_values = {
                            header[j]: (row[j] if j < len(row) else None)
                            for j in range(len(header))
                        }
                        output.append(f"Row {i}: {json.dumps(row_values, default=str, indent=2)}")
                return "\n\n".join(output), {"doc_type": (detected_doc_type or None)}
            except ImportError as e:
                logger.error(f"Missing dependency for XLSX parsing: {e}")
                raise

        if ft == "docx":
            try:
                from docx import Document as DocxDocument

                doc = DocxDocument(
                    BytesIO(raw if isinstance(raw, bytes) else (raw_text or "").encode("utf-8"))
                )
                paragraphs = [p.text for p in doc.paragraphs if p.text]
                return "\n\n".join(paragraphs), {}
            except ImportError as e:
                logger.error(f"Missing dependency for DOCX parsing: {e}")
                raise

        if ft == "pdf":
            try:
                from pypdf import PdfReader

                reader = PdfReader(
                    BytesIO(raw if isinstance(raw, bytes) else (raw_text or "").encode("utf-8"))
                )
                if getattr(reader, "is_encrypted", False):
                    try:
                        reader.decrypt("")
                    except Exception:
                        raise ValueError("Encrypted PDF is not supported without a password.")
                page_blocks: list[str] = []
                for i, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text() or ""
                    except Exception:
                        page_text = ""
                    page_blocks.append(f"[Page {i+1}]\n{page_text}")
                return "\n\n".join(page_blocks), {}
            except ImportError as e:
                logger.error(f"Missing dependency for PDF parsing: {e}")
                raise

        # Fallback: best-effort decode
        return (
            raw_text or (raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else "")
        ), {}
    except Exception as exc:
        logger.error(f"Error extracting text for file_type={file_type}: {exc}", exc_info=True)
        return (raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else (raw or "")), {}


def save_file(file_id, filename, file_type, content, detected_doc_type=None):
    """
    Save a file to the database

    Args:
        file_id: Unique identifier for the file
        filename: Name of the file
        file_type: File type/extension
        content: File content

    Returns:
        The document object
    """
    try:
        # Determine document type based on detection or filename heuristics
        doc_type = (detected_doc_type or "").lower() if detected_doc_type else "other"
        if doc_type == "other":
            if (
                "course" in filename.lower()
                or "lesson" in filename.lower()
                or "university" in filename.lower()
            ):
                doc_type = "course"
            elif (
                "documentation" in filename.lower()
                or "docs" in filename.lower()
                or "manual" in filename.lower()
            ):
                doc_type = "documentation"

        # Create a document object
        document = Document(
            id=file_id,
            filename=filename,
            file_type=file_type,
            content=content,
            file_metadata={"source": "user_upload", "file_type": file_type, "doc_type": doc_type},
        )

        # Save to database
        db.session.add(document)
        db.session.commit()

        logger.debug(
            f"File saved to database: {filename} (ID: {file_id}, doc_type={doc_type}, preview='{_safe_preview(content)}')"
        )
        return document

    except Exception as e:
        db.session.rollback()
        logger.error(f"Error saving file: {str(e)}", exc_info=True)
        raise


def get_all_files():
    """
    Get all files metadata (without content)

    Returns:
        List of file metadata dictionaries
    """
    try:
        documents = Document.query.all()
        return [doc.to_dict() for doc in documents]
    except Exception as e:
        logger.error(f"Error getting all files: {str(e)}", exc_info=True)
        return []


def get_file_by_id(file_id):
    """
    Get a file by ID

    Args:
        file_id: ID of the file to retrieve

    Returns:
        Tuple of (file_metadata, file_content) or (None, None) if not found
    """
    try:
        document = db.session.get(Document, file_id)
        if not document:
            return None, None

        return document.to_dict(), document.content
    except Exception as e:
        logger.error(f"Error getting file by ID: {str(e)}", exc_info=True)
        return None, None


def delete_file_by_id(file_id):
    """
    Delete a file by ID

    Args:
        file_id: ID of the file to delete

    Returns:
        Boolean indicating success
    """
    try:
        document = db.session.get(Document, file_id)
        if not document:
            return False

        # Get the chunks associated with this file
        chunk_ids = [chunk.vector_id for chunk in document.chunks]

        # Delete the chunks from Pinecone
        for chunk_id in chunk_ids:
            delete_from_pinecone(chunk_id)

        # Delete the document from the database
        db.session.delete(document)
        db.session.commit()

        logger.debug(f"File deleted from database: {file_id}")
        return True

    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting file: {str(e)}", exc_info=True)
        return False


def scan_knowledge_base():
    """
    Retrieve a list of all document chunks from Pinecone for display purposes.
    This function is read-only and doesn't modify the database.

    Returns:
        Dictionary with pinecone_stats and pinecone_documents
    """
    logger.info("Retrieving document information from Pinecone")

    # Get all documents from Pinecone
    pinecone_data = list_all_vectors()

    # Return statistics and document list
    return {
        "pinecone_stats": pinecone_data["stats"],
        "pinecone_documents": pinecone_data["documents"],
    }


def process_file_for_rag(file_id, filename, content, file_type):
    """
    Process a file for RAG by chunking it and upserting to Pinecone

    Args:
        file_id: ID of the file
        filename: Name of the file
        content: Content of the file
        file_type: Type of the file

    Returns:
        List of chunk IDs
    """
    try:
        # 'content' should already be normalized text when called from routes
        parsed_content = content or ""
        chunks = chunk_text(parsed_content)

        # Track chunk IDs
        chunk_ids = []

        # Find the document
        document = db.session.get(Document, file_id)
        if not document:
            logger.error(f"Document not found for chunking: {file_id}")
            return []

        for i, chunk_content in enumerate(chunks):
            chunk_id = f"{file_id}_chunk_{i}"
            document = Document.query.get(file_id)
            if not document:
                logger.error(f"Document not found for metadata: {file_id}")
                continue
            doc_type = (document.file_metadata or {}).get("doc_type", "other")

            # Metadata for the chunk
            metadata = {
                "content": chunk_content,
                "source_file_id": file_id,
                "source_filename": filename,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "title": filename,
                "url": f"/api/files/{file_id}",
                "type": doc_type,
            }

            # Upsert to Pinecone
            success = upsert_to_pinecone(chunk_id, chunk_content, metadata)

            if success:
                # Create a document chunk record
                document_chunk = DocumentChunk(
                    id=chunk_id,
                    document_id=file_id,
                    chunk_index=i,
                    content=chunk_content,
                    vector_id=chunk_id,
                )
                db.session.add(document_chunk)
                chunk_ids.append(chunk_id)

        # Commit the chunks to the database
        db.session.commit()

        logger.debug(f"File processed for RAG: {filename} (ID: {file_id}, Chunks: {len(chunks)})")
        return chunk_ids

    except Exception as e:
        db.session.rollback()
        logger.error(f"Error processing file for RAG: {str(e)}", exc_info=True)
        return []
