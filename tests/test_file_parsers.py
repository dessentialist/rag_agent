import json
from io import BytesIO

import pytest

from services.file_service import extract_text_and_metadata


def test_txt_parser_simple():
    text = "Hello world\nThis is a test."
    parsed, meta = extract_text_and_metadata("txt", text)
    assert "Hello world" in parsed
    assert meta == {}


def test_md_parser_simple():
    md = "# Title\n\nSome **bold** text."
    parsed, meta = extract_text_and_metadata("md", md)
    assert "Title" in parsed
    assert meta == {}


def test_json_parser_detects_doc_type():
    obj = {"doc_type": "documentation", "content": "stuff"}
    parsed, meta = extract_text_and_metadata("json", json.dumps(obj))
    assert "documentation" in parsed
    assert meta.get("doc_type") == "documentation"


def test_jsonl_parser_detects_doc_type_from_first_record():
    records = [
        json.dumps({"Category": "course", "x": 1}),
        json.dumps({"y": 2}),
    ]
    parsed, meta = extract_text_and_metadata("jsonl", "\n".join(records))
    assert "Record 1" in parsed
    assert meta.get("doc_type") == "course"


def test_csv_parser_pretty_print_and_doc_type():
    csv_text = "Category,Name\ncourse,Intro\ncourse,Advanced"
    parsed, meta = extract_text_and_metadata("csv", csv_text)
    assert "CSV Headers" in parsed
    assert "Row 1" in parsed
    assert meta.get("doc_type") == "course"


def test_xlsx_parser_roundtrip(tmp_path):
    try:
        from openpyxl import Workbook
    except Exception:
        pytest.skip("openpyxl not available")

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["doc_type", "Name"])  # header
    ws.append(["documentation", "Intro"])  # first row with doc_type
    ws.append(["documentation", "Advanced"])  # second row
    xlsx_path = tmp_path / "sample.xlsx"
    wb.save(xlsx_path)

    parsed, meta = extract_text_and_metadata("xlsx", xlsx_path.read_bytes())
    assert "XLSX Headers" in parsed
    assert meta.get("doc_type") == "documentation"


def test_docx_parser_paragraphs(tmp_path):
    try:
        from docx import Document as DocxDocument
    except Exception:
        pytest.skip("python-docx not available")

    doc = DocxDocument()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    docx_path = tmp_path / "sample.docx"
    doc.save(docx_path)

    parsed, meta = extract_text_and_metadata("docx", docx_path.read_bytes())
    assert "First paragraph" in parsed
    assert meta == {}


def test_pdf_parser_pages(tmp_path):
    # Generate a simple PDF using pypdf if possible
    try:
        from pypdf import PdfWriter
    except Exception:
        pytest.skip("pypdf not available")

    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    pdf_path = tmp_path / "sample.pdf"
    with open(pdf_path, "wb") as f:
        writer.write(f)

    parsed, meta = extract_text_and_metadata("pdf", pdf_path.read_bytes())
    # Even blank pages produce a [Page 1] marker
    assert "[Page 1]" in parsed
    assert meta == {}


